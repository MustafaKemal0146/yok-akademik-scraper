from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
from tqdm import tqdm
from colorama import Fore, Style, init
import time
import pandas as pd
from fpdf import FPDF
from fpdf.fonts import FontFace
from datetime import datetime
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tqdm import tqdm
import re
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager

class YOKScraper:
    def __init__(self, url):
        """
        YÖK Akademik Scraper sınıfını başlatır
        Args:
            url: Akademisyenin YÖK profil URL'i
        """
        if not url.startswith("https://akademik.yok.gov.tr/"):
            raise ValueError("Geçersiz YÖK Akademik URL'i")
        self.url = url
        self.driver = None
        self.wait = None
        
    def __enter__(self):
        self.setup_driver()
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.driver:
            self.driver.quit()

    def setup_driver(self):
        """Selenium webdriver'ı optimize edilmiş ayarlarla başlatır"""
        options = webdriver.ChromeOptions()
        options.add_argument('--headless=new')
        options.add_argument('--disable-gpu')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--window-size=1920,1080')
        
        # Performans optimizasyonları
        options.add_argument('--disable-extensions')
        options.add_argument('--disable-infobars')
        options.add_argument('--disable-notifications')
        options.add_argument('--disable-logging')
        options.add_argument('--disable-web-security')
        options.add_argument('--dns-prefetch-disable')
        options.page_load_strategy = 'eager'  # Sayfa tam yüklenmeden işlem yapmaya başla
        
        prefs = {
            'profile.managed_default_content_settings.images': 2,  # Resimleri devre dışı bırak
            'disk-cache-size': 4096,  # Cache boyutunu artır
            'profile.default_content_setting_values.notifications': 2
        }
        options.add_experimental_option('prefs', prefs)
        
        service = Service(ChromeDriverManager().install())
        self.driver = webdriver.Chrome(service=service, options=options)
        self.wait = WebDriverWait(self.driver, 5)  # Timeout süresini düşür

    def get_academic_info(self):
        """
        Akademisyenin görev ve öğrenim bilgilerini çeker
        Returns:
            dict: Akademik bilgileri içeren sözlük
        """
        try:
            self.driver.get(self.url)
            time.sleep(3)
            
            academic_info = {
                'duties': [],
                'education': []
            }
            
            # Akademik Görevler
            duties = self.driver.find_elements(By.CSS_SELECTOR, "div.col-md-6:nth-child(1) ul.timeline li")
            current_year = None
            
            for duty in duties:
                try:
                    # Başlık kontrolü
                    if "Akademik Görevler" in duty.text:
                        continue
                        
                    # Yıl etiketi
                    year = duty.find_elements(By.CSS_SELECTOR, "span.bg-light-blue")
                    if year:
                        current_year = year[0].text.strip()
                        continue
                    
                    # Görev bilgileri
                    position = duty.find_elements(By.CSS_SELECTOR, "div.timeline-footer a.btn-success")
                    institution = duty.find_elements(By.CSS_SELECTOR, "div.timeline-item h4")
                    department = duty.find_elements(By.CSS_SELECTOR, "div.timeline-item h5")
                    
                    if position and institution:
                        duty_text = f"{current_year} - {position[0].text.strip()} - {institution[0].text.strip()}"
                        if department:
                            duty_text += f" - {department[0].text.strip()}"
                        academic_info['duties'].append(duty_text)
                except Exception as e:
                    continue
            
            # Öğrenim Bilgileri
            education = self.driver.find_elements(By.CSS_SELECTOR, "div.col-md-6:nth-child(2) ul.timeline li")
            current_year = None
            
            for edu in education:
                try:
                    # Başlık kontrolü
                    if "Öğrenim Bilgisi" in edu.text:
                        continue
                        
                    # Yıl etiketi
                    year = edu.find_elements(By.CSS_SELECTOR, "span.bg-light-blue")
                    if year:
                        current_year = year[0].text.strip()
                        continue
                    
                    # Öğrenim bilgileri
                    degree = edu.find_elements(By.CSS_SELECTOR, "div.timeline-footer a.btn-info")
                    institution = edu.find_elements(By.CSS_SELECTOR, "div.timeline-item h4")
                    department = edu.find_elements(By.CSS_SELECTOR, "div.timeline-item h5")
                    thesis = edu.find_elements(By.CSS_SELECTOR, "div.timeline-item h6")
                    
                    if degree and institution:
                        edu_text = f"{current_year} - {degree[0].text.strip()} - {institution[0].text.strip()}"
                        if department:
                            edu_text += f" - {department[0].text.strip()}"
                        if thesis:
                            edu_text += f" - {thesis[0].text.strip()}"
                        academic_info['education'].append(edu_text)
                except Exception as e:
                    continue
                
            return academic_info
        except Exception as e:
            print(f"Akademik bilgiler çekilirken hata: {e}")
            return None

    def get_books(self):
        """
        Akademisyenin kitaplarını çeker
        Returns:
            list: Kitap bilgilerini içeren liste
        """
        try:
            books_tab = self.wait.until(EC.element_to_be_clickable(
                (By.CSS_SELECTOR, "li#booksMenu a")))
            books_tab.click()
            
            # Sabit bekleme yerine dinamik bekleme
            self.wait.until(EC.presence_of_element_located(
                (By.CSS_SELECTOR, "div.container-fluid div.row div")))
            
            books = []
            
            # Ana kitap container'ı
            book_container = self.driver.find_element(
                By.CSS_SELECTOR, 
                "div.col-lg-10.col-md-10.col-sm-12.col-xs-12 > div.container-fluid > div.row > div"
            )
            
            # Kitap elementleri - doğrudan div.projects altındaki div'leri al
            book_elements = book_container.find_elements(By.CSS_SELECTOR, "div.projects > div")
            
            pbar = tqdm(book_elements, desc="Kitaplar çekiliyor", unit="kitap")
            
            for book in pbar:
                try:
                    # Başlık - strong elementi içinde
                    title = book.find_element(By.CSS_SELECTOR, "strong").text.strip()
                    
                    # Detay bilgileri - p elementi içinde
                    info_text = book.find_element(By.CSS_SELECTOR, "p").text.strip()
                    
                    # Yazarlar - virgülle ayrılmış ilk kısım
                    authors = info_text.split("Yayın Yeri:")[0].strip()
                    
                    # Diğer bilgiler
                    publisher = ""
                    edition = ""
                    pages = ""
                    isbn = ""
                    year = ""
                    
                    # Bilgileri parçala
                    info_parts = info_text.split(",")
                    for part in info_parts:
                        part = part.strip()
                        if "Yayın Yeri:" in part:
                            publisher = part.split("Yayın Yeri:")[1].strip()
                        elif "Basım sayısı:" in part:
                            edition = part.split("Basım sayısı:")[1].strip()
                        elif "Sayfa sayısı:" in part:
                            pages = part.split("Sayfa sayısı:")[1].strip()
                        elif "ISBN:" in part:
                            isbn = part.split("ISBN:")[1].strip()
                    
                    # Yıl bilgisi için tüm metni kontrol et
                    year_match = re.search(r'\b(19|20)\d{2}\b', info_text)
                    if year_match:
                        year = year_match.group()
                    else:
                        # Yıl bilgisi için label elementlerini kontrol et
                        year_elements = book.find_elements(By.CSS_SELECTOR, "span.label.label-info")
                        if year_elements:
                            year = year_elements[0].text.strip()
                    
                    book_info = {
                        'title': title,
                        'authors': authors,
                        'publisher': publisher,
                        'year': year,
                        'edition': edition,
                        'pages': pages,
                        'isbn': isbn
                    }
                    books.append(book_info)
                    
                    pbar.set_description(f"Kitap işleniyor: {title[:30]}...")
                    
                except Exception as e:
                    print(f"Kitap bilgisi çekilirken hata: {e}")
                    continue
            
            return books
            
        except Exception as e:
            print(f"Kitaplar çekilirken hata: {e}")
            return []

    def get_articles(self):
        try:
            articles_tab = self.wait.until(EC.element_to_be_clickable(
                (By.CSS_SELECTOR, "li#articlesMenu a")))
            articles_tab.click()
            
            # Dinamik bekleme
            try:
                self.wait.until(EC.presence_of_element_located(
                    (By.XPATH, "//div[@id='articles']//table/tbody/tr")))
            except:
                print("Makale verisi bulunamadı veya yüklenmedi")
                return []
            
            articles = []
            
            try:
                # Tüm tr elementlerini bul - daha spesifik xpath
                rows = self.driver.find_elements(
                    By.XPATH, 
                    "//div[@id='articles']//table/tbody/tr[td[2]]"  # En az 2 td'si olan tr'ler
                )
                
                if not rows:  # Eğer rows boşsa
                    print("Makale satırları bulunamadı")
                    return []
                    
                pbar = tqdm(rows, desc="Makaleler çekiliyor", unit="makale")
                
                for row in pbar:
                    try:
                        cells = row.find_elements(By.TAG_NAME, "td")
                        if not cells or len(cells) < 2:  # None ve uzunluk kontrolü
                            continue
                        
                        info_cell = cells[1]  # İkinci hücre
                        if not info_cell:  # info_cell None kontrolü
                            continue
                        
                        # Başlık kontrolü
                        title_elements = info_cell.find_elements(By.TAG_NAME, "strong")
                        if not title_elements:
                            continue
                        title = title_elements[0].text.strip()
                        
                        # Tüm metin içeriği
                        full_text = info_cell.text.strip()
                        if not full_text:  # Boş metin kontrolü
                            continue
                        
                        # Yazarlar
                        try:
                            text_after_title = full_text.split(title)[1].strip()
                            authors = text_after_title.split("Yayın Yeri:")[0].strip().strip(',')
                        except:
                            authors = ""
                        
                        # Yayın yeri ve yıl
                        try:
                            if "Yayın Yeri:" in full_text:
                                pub_text = full_text.split("Yayın Yeri:")[1]
                                pub_match = re.search(r'(.*?)(?=,\s*(?:19|20)\d{2}|$|\n|<p>)', pub_text)
                                if pub_match:
                                    publication_place = pub_match.group(1).strip()
                                else:
                                    publication_place = pub_text.split(',')[0].strip()
                            else:
                                publication_place = ""
                            
                            year_match = re.search(r'\b(19|20)\d{2}\b', full_text)
                            year = year_match.group() if year_match else ""
                            
                        except Exception as e:
                            print(f"Yayın yeri/yıl ayrıştırma hatası: {e}")
                            publication_place = ""
                            year = ""
                        
                        # Etiketler - üçüncü hücre
                        labels = []
                        if len(cells) > 2:
                            label_elements = cells[2].find_elements(By.TAG_NAME, "span")
                            labels = [label.text.strip() for label in label_elements if label.text.strip()]
                        
                        # DOI - dördüncü hücre
                        doi = ""
                        if len(cells) > 3:
                            doi_links = cells[3].find_elements(By.TAG_NAME, "a")
                            for link in doi_links:
                                href = link.get_attribute("href")
                                if href and "doi.org" in href:
                                    doi = href
                                    break
                        
                        article_info = {
                            'title': title,
                            'authors': authors,
                            'publication_place': publication_place,
                            'year': year,
                            'labels': labels,
                            'doi': doi
                        }
                        articles.append(article_info)
                        
                        pbar.set_description(f"Makale işleniyor: {title[:30]}...")
                        
                    except Exception as e:
                        print(f"Makale bilgisi çekilirken hata: {e}")
                        continue
                
                return articles
                
            except Exception as e:
                print(f"Makale listesi işlenirken hata: {e}")
                return []
                
        except Exception as e:
            print(f"Makaleler çekilirken hata: {e}")
            return []

    def get_proceedings(self):
        """
        Akademisyenin bildirilerini çeker
        Returns:
            list: Bildiri bilgilerini içeren liste
        """
        try:
            proceedings_tab = self.wait.until(EC.element_to_be_clickable(
                (By.CSS_SELECTOR, "li#proceedingMenu a")))
            proceedings_tab.click()
            
            # Bildiri elementlerinin yüklenmesini bekle
            self.wait.until(EC.presence_of_element_located(
                (By.XPATH, '(//table)[1]/tbody/tr')))
            
            proceedings = []
            rows = self.driver.find_elements(By.XPATH, '(//table)[1]/tbody/tr')
            
            pbar = tqdm(rows, desc="Bildiriler çekiliyor", unit="bildiri")
            
            for proceeding in pbar:
                try:
                    # İkinci hücreyi al (index 1)
                    second_td = proceeding.find_elements(By.XPATH, './/td')[1]
                    outer_html = second_td.get_attribute('outerHTML')
                    
                    # Başlık
                    title = proceeding.find_element(By.CSS_SELECTOR, "strong").text.strip()
                    
                    # Yazarlar - HTML parsing ile
                    try:
                        # <p></p> ile <p> arasındaki veriyi al
                        p_content = outer_html.split('<p></p>')[1].split('<p>')[0].strip()
                        parts = p_content.replace('\n, ', ' , ').split(' , ')
                        p_authors_html = parts[0].strip() if len(parts) > 0 else ''
                        
                        # HTML etiketlerini kaldır ve yazarları ayır
                        soup = BeautifulSoup(p_authors_html, 'html.parser')
                        authors_list = [author.strip() for author in soup.get_text().split(',')]
                        
                        # Son yazar adında bulunan tarih bilgisini ayır
                        if authors_list:
                            authors_list[-1] = authors_list[-1].split('\n ')[0].strip()
                        
                        authors = ', '.join(filter(None, authors_list))
                    except Exception as e:
                        print(f"Yazar ayrıştırma hatası: {e}")
                        authors = ""
                    
                    # Yayın yeri ve yıl
                    try:
                        if "Yayın Yeri:" in outer_html:
                            pub_text = outer_html.split("Yayın Yeri:")[1]
                            pub_match = re.search(r'(.*?)(?=,\s*(?:19|20)\d{2}|$|\n|<p>)', pub_text)
                            if pub_match:
                                publication_info = pub_match.group(1).strip()
                                publication_info = re.sub(r'[,\s]+$', '', publication_info)
                                publication_info = publication_info.strip(' :,')
                            else:
                                publication_info = ""
                        else:
                            publication_info = ""
                        
                        year_match = re.search(r'\b(19|20)\d{2}\b', outer_html)
                        year = year_match.group() if year_match else ""
                        
                    except Exception as e:
                        print(f"Yayın yeri/yıl ayrıştırma hatası: {e}")
                        publication_info = ""
                        year = ""
                    
                    # Tür bilgisi
                    try:
                        span_tags = second_td.find_elements(By.XPATH, './/p[2]/span')
                        types = []
                        if len(span_tags) > 0:
                            scope = span_tags[0].text.strip()
                            if scope:
                                types.append(scope)
                        if len(span_tags) > 1:
                            proc_type = span_tags[1].text.strip()
                            if proc_type:
                                types.append(proc_type)
                        proceeding_type = " | ".join(types)
                    except Exception as e:
                        print(f"Tür bilgisi çekme hatası: {e}")
                        proceeding_type = ""
                    
                    proceeding_info = {
                        'title': title,
                        'authors': authors,
                        'publication_place': publication_info,
                        'year': year,
                        'type': proceeding_type
                    }
                    proceedings.append(proceeding_info)
                    
                    pbar.set_description(f"Bildiri işleniyor: {title[:30]}...")
                    
                except Exception as e:
                    print(f"Bildiri bilgisi çekilirken hata: {e}")
                    continue
            
            return proceedings
            
        except Exception as e:
            print(f"Bildiriler çekilirken hata: {e}")
            return []

    def get_presentations(self, html_content):
        presentations = []
        soup = BeautifulSoup(html_content, 'html.parser')
        container = soup.find('div', class_='container-fluid')
        if not container:
            print("Container for presentations not found.")
            return presentations

        all_tab = container.find('div', id='all')
        if not all_tab:
            print("Tab for all presentations not found.")
            return presentations

        rows = all_tab.find_all('tr')
        for row in rows:
            try:
                title = row.find('span', class_='baslika').find('strong').text.strip()
            except Exception as e:
                print(f"Title not found: {e}")
                continue

            try:
                authors = ', '.join([a.text.strip() for a in row.find_all('a', class_='popoverData')])
            except Exception as e:
                print(f"Authors not found: {e}")
                authors = ''

            try:
                pub_info = row.find('span', class_='baslika').find_next_sibling(text=True).strip()
            except Exception as e:
                print(f"Publication info not found: {e}")
                pub_info = ''

            try:
                labels = ', '.join([label.text.strip() for label in row.find_all('span', class_='label')])
            except Exception as e:
                print(f"Labels not found: {e}")
                labels = ''

            presentations.append({
                'title': title,
                'authors': authors,
                'publication_info': pub_info,
                'labels': labels
            })

        return presentations

    def get_articles_from_html(self, html_content):
        articles = []
        soup = BeautifulSoup(html_content, 'html.parser')
        tbody = soup.find('tbody', class_='searchable')
        if not tbody:
            print("Table body for articles not found.")
            return articles

        rows = tbody.find_all('tr')
        for row in rows:
            try:
                title_element = row.find('span', class_='baslika').find('strong').find('a')
                title = title_element.text.strip()
                link = title_element['href']
            except Exception as e:
                print(f"Title or link not found: {e}")
                continue

            try:
                authors = ', '.join([a.text.strip() for a in row.find_all('a', class_='popoverData')])
            except Exception as e:
                print(f"Authors not found: {e}")
                authors = ''

            try:
                pub_info = row.find('span', class_='baslika').find_next_sibling(text=True).strip()
                pub_parts = pub_info.split(',')
                publication_place = pub_parts[0].replace('Yayın Yeri:', '').strip()
                year = pub_parts[1].strip()
            except Exception as e:
                print(f"Publication info not found: {e}")
                publication_place = ''
                year = ''

            try:
                labels = ', '.join([label.text.strip() for label in row.find_all('span', class_='label')])
            except Exception as e:
                print(f"Labels not found: {e}")
                labels = ''

            try:
                doi_link = row.find('a', href=True, target='_blank')['href']
            except Exception as e:
                print(f"DOI link not found: {e}")
                doi_link = ''

            articles.append({
                'title': title,
                'link': link,
                'authors': authors,
                'publication_place': publication_place,
                'year': year,
                'labels': labels,
                'doi_link': doi_link
            })

        return articles

    def get_books_from_html(self, html_content):
        books = []
        soup = BeautifulSoup(html_content, 'html.parser')
        projects_div = soup.find('div', class_='projects')
        if not projects_div:
            print("Projects container for books not found.")
            return books

        rows = projects_div.find_all('div', class_='row')
        for row in rows:
            try:
                title = row.find('strong').text.strip()
            except Exception as e:
                print(f"Title not found: {e}")
                continue

            try:
                details = row.find('p').text.strip()
                details_parts = details.split(',')
                authors = details_parts[0].strip()
                publication_place = details_parts[1].replace('Yayın Yeri:', '').strip()
                edition = details_parts[2].replace('Basım sayısı:', '').strip()
                page_count = details_parts[3].replace('Sayfa sayısı:', '').strip()
                isbn = details_parts[4].replace('ISBN:', '').strip()
            except Exception as e:
                print(f"Details not found: {e}")
                authors = publication_place = edition = page_count = isbn = ''

            try:
                year = row.find('span', class_='label-info').text.strip()
            except Exception as e:
                print(f"Year not found: {e}")
                year = ''

            try:
                book_type = row.find('span', class_='label-primary').text.strip()
            except Exception as e:
                print(f"Book type not found: {e}")
                book_type = ''

            books.append({
                'title': title,
                'authors': authors,
                'publication_place': publication_place,
                'edition': edition,
                'page_count': page_count,
                'isbn': isbn,
                'year': year,
                'type': book_type
            })

        return books

    def fetch_articles(self):
        """Makaleleri çeker"""
        try:
            wait = WebDriverWait(self.driver, 10)
            
            # "Makaleler" linkine tıklayın
            articles_link = wait.until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, 'li#articleMenu a'))
            )
            articles_link.click()

            # İlk tablonun yüklenmesini bekleyin
            wait.until(
                EC.presence_of_element_located((By.XPATH, '(//table)[1]/tbody'))
            )
            
            # İlk tablodaki makaleleri çekin
            articles = []
            rows = self.driver.find_elements(By.XPATH, '(//table)[1]/tbody/tr')
            pbar = tqdm(rows, desc="Makaleler işleniyor", unit="makale", colour="green")
            
            for idx, row in enumerate(pbar):
                try:
                    # Makale numarası
                    first_td = row.find_elements(By.XPATH, './/td')[0]
                    p_no = first_td.text.strip() if first_td else 'N/A'
                    
                    # Makale detayları
                    second_td = row.find_elements(By.XPATH, './/td')[1]
                    outer_html = second_td.get_attribute('outerHTML')
                    title_tag = second_td.find_element(By.XPATH, './/a[@data-toggle="modal"]')
                    p_title = title_tag.text.strip()

                    # İçerik ayrıştırma
                    p_content = outer_html.split('<p></p>')[1].split('<p>')[0].strip()
                    parts = p_content.replace('\n, ', ' , ').split(' , ')
                    p_authors_html = parts[0].strip() if len(parts) > 0 else 'N/A'
                    p_event = parts[1].strip().replace('Yayın Yeri:', '').strip() if len(parts) > 1 else 'N/A'
                    p_year_info = parts[2].strip() if len(parts) > 2 else 'N/A'

                    # Yazarları ayır
                    soup = BeautifulSoup(p_authors_html, 'html.parser')
                    p_authors = [author.strip() for author in soup.get_text().split(',')]

                    # Etiketleri al
                    span_tags = second_td.find_elements(By.XPATH, './/p[2]/span')
                    p_scope = span_tags[0].text.strip() if len(span_tags) > 0 else 'N/A'
                    p_type = span_tags[1].text.strip() if len(span_tags) > 1 else 'N/A'
                    p_index = span_tags[2].text.strip() if len(span_tags) > 2 else 'N/A'
                    p_article_type = span_tags[3].text.strip() if len(span_tags) > 3 else 'N/A'
                    
                    # DOI bilgisi
                    doi_tags = second_td.find_elements(By.XPATH, './/p[2]/a')
                    p_doi = doi_tags[0].get_attribute('href') if doi_tags else 'N/A'

                    # Makale bilgilerini ekle
                    articles.append({
                        'no': p_no,
                        'title': p_title,
                        'authors': p_authors,
                        'publication': p_event,
                        'year': p_year_info,
                        'scope': p_scope,
                        'type': p_type,
                        'index': p_index,
                        'article_type': p_article_type,
                        'doi': p_doi
                    })
                    pbar.set_description(f"Makale işleniyor: {p_title[:30]}...")
                    
                except Exception as e:
                    print(Fore.RED + f"HATA: Makale çekilemedi! (Article {idx + 1:2d}): {e}")

            return articles
            
        except Exception as e:
            print(Fore.RED + f"Bir hata oluştu: {e}")
            return []

    def scrape_all(self):
        try:
            self.setup_driver()
            total_steps = 7  # Dersler için +1
            with tqdm(total=total_steps, desc="Veriler çekiliyor", colour="green") as pbar:
                try:
                    pbar.set_description("Akademik bilgiler çekiliyor")
                    academic_info = self.get_academic_info()
                    pbar.update(1)
                    
                    pbar.set_description("Kitaplar çekiliyor")
                    books = self.get_books()
                    pbar.update(1)
                    
                    pbar.set_description("Makaleler çekiliyor")
                    articles = self.fetch_articles()
                    pbar.update(1)
                    
                    pbar.set_description("Bildiriler çekiliyor")
                    proceedings = self.get_proceedings()
                    pbar.update(1)
                    
                    pbar.set_description("Projeler çekiliyor")
                    projects = self.get_projects()
                    pbar.update(1)
                    
                    pbar.set_description("Dersler çekiliyor")
                    lessons = self.get_lessons()
                    pbar.update(1)
                    
                    results = {
                        'academic_info': academic_info,
                        'books': books,
                        'articles': articles,
                        'proceedings': proceedings,
                        'projects': projects,
                        'lessons': lessons
                    }
                    
                    pbar.set_description("İşlem tamamlandı!")
                    return results
                    
                except Exception as e:
                    print(Fore.RED + f"Veri çekme hatası: {e}")
                    pbar.update(total_steps - pbar.n)
                    return None
        finally:
            if self.driver:
                self.driver.quit()

    def save_to_word(self, data, filename='academic_info.docx'):
        """
        Verileri Word dosyasına belirtilen sırada kaydeder
        Args:
            data: Kaydedilecek veriler
            filename: Kaydedilecek dosya adı
        """
        if not data:
            print("Kaydedilecek veri bulunamadı!")
            return
        
        if not isinstance(data, dict):
            print("Geçersiz veri formatı!")
            return
        
        doc = Document()
        
        # Başlık stilini ayarla
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # 1. Öğrenim Bilgileri
        if 'academic_info' in data and data['academic_info']:
            doc.add_heading('Öğrenim Bilgileri', level=1)
            if 'education' in data['academic_info']:
                for edu in data['academic_info']['education']:
                    # Öğrenim bilgisini parçalara ayır
                    parts = edu.split(' - ')
                    if len(parts) >= 4:
                        year = parts[0]
                        degree = parts[1]
                        institution = parts[2]
                        department = parts[3]
                        thesis = parts[4] if len(parts) > 4 else None

                        p = doc.add_paragraph()
                        p.add_run(f"{institution}").bold = True
                        p.add_run(f"\n{department}")
                        p.add_run(f"\n{degree}")
                        p.add_run(f"\n{year}")
                        if thesis:
                            p.add_run(f"\n{thesis}")
                        doc.add_paragraph()  # Boş satır ekle
                    else:
                        # Parçalama başarısız olursa orijinal formatı kullan
                        doc.add_paragraph(edu)
            doc.add_page_break()
        
        # 2. Akademik Görevler
        if 'academic_info' in data and data['academic_info']:
            doc.add_heading('Akademik Görevler', level=1)
            if 'duties' in data['academic_info']:
                for duty in data['academic_info']['duties']:
                    doc.add_paragraph(duty)
            doc.add_page_break()
        
        # 3. Kitaplar
        if 'books' in data and data['books']:
            doc.add_heading('Kitaplar', level=1)
            for book in data['books']:
                p = doc.add_paragraph()
                p.add_run(f"Başlık: {book['title']}").bold = True
                p.add_run(f"\nYazarlar: {book['authors']}")
                p.add_run(f"\nYayınevi: {book['publisher']}")
                if book['year']:
                    p.add_run(f"\nYıl: {book['year']}")
                if book['edition']:
                    p.add_run(f"\nBasım sayısı: {book['edition']}")
                if book['pages']:
                    p.add_run(f"\nSayfa sayısı: {book['pages']}")
                if book['isbn']:
                    p.add_run(f"\nISBN: {book['isbn']}")
                doc.add_paragraph()
            doc.add_page_break()
        
        # 4. Makaleler
        if 'articles' in data and data['articles']:
            doc.add_heading('Makaleler', level=1)
            for article in data['articles']:
                p = doc.add_paragraph()
                p.add_run(f"Başlık: {article['title']}").bold = True
                p.add_run(f"\nYazarlar: {', '.join(article['authors'])}")
                p.add_run(f"\nYayın Yeri: {article['publication']}")
                p.add_run(f"\nYıl: {article['year']}")
                p.add_run(f"\nKapsam: {article['scope']}")
                p.add_run(f"\nTür: {article['type']}")
                p.add_run(f"\nİndeks: {article['index']}")
                p.add_run(f"\nMakale Türü: {article['article_type']}")
                if article['doi'] != 'N/A':
                    p.add_run(f"\nDOI: {article['doi']}")
                doc.add_paragraph()
            doc.add_page_break()
        
        # 5. Bildiriler
        if 'proceedings' in data and data['proceedings']:
            doc.add_heading('Bildiriler', level=1)
            for proc in data['proceedings']:
                p = doc.add_paragraph()
                p.add_run(f"Başlık: {proc['title']}").bold = True
                p.add_run(f"\nYazarlar: {proc['authors']}")
                p.add_run(f"\nYayın Yeri: {proc['publication_place']}")
                p.add_run(f"\nYıl: {proc['year']}")
                
                # Tür bilgisini ayır
                if '|' in proc['type']:
                    proceeding_type, presentation_type = [t.strip() for t in proc['type'].split('|')]
                    p.add_run(f"\nTür: {proceeding_type}")
                    p.add_run(f"\nSunum Şekli: {presentation_type}")
                else:
                    p.add_run(f"\nTür: {proc['type']}")
                    
                doc.add_paragraph()
        
        # Projeler
        if 'projects' in data and data['projects']:
            doc.add_heading('Projeler', level=1)
            for project in data['projects']:
                p = doc.add_paragraph()
                p.add_run(f"Proje Adı: {project['title']}").bold = True
                p.add_run(f"\nKatkısı Geçenler: {project['contributors']}")
                p.add_run(f"\nProjenin Başlangıç ve Bitiş Tarihi: {project['dates']}")
                p.add_run(f"\nProjenin Yürütüldüğü Yer: {project['institution']}")
                p.add_run(f"\nSüreç: {project['status']}")
                p.add_run(f"\nKullanılan Bütçe: {project['budget']}")
                doc.add_paragraph()
            doc.add_page_break()
        
        # Dersler
        if 'lessons' in data and data['lessons']:
            doc.add_heading('Verdiği Dersler', level=1)
            
            for level, info in data['lessons'].items():
                if info['lessons']:  # Eğer dersler varsa
                    doc.add_heading(info['name'], level=2)
                    
                    # Tablo oluştur
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    
                    # Başlık satırı
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'Dönem'
                    header_cells[1].text = 'Ders Adı'
                    header_cells[2].text = 'Dili'
                    header_cells[3].text = 'Saat'
                    
                    # Ders bilgilerini ekle
                    for lesson in info['lessons']:
                        row_cells = table.add_row().cells
                        row_cells[0].text = lesson['term']
                        row_cells[1].text = lesson['name']
                        row_cells[2].text = lesson['language']
                        row_cells[3].text = lesson['hours']
                    
                    doc.add_paragraph()  # Boşluk ekle
            doc.add_page_break()
        
        # Dosyayı kaydet
        doc.save(filename)
        print(f"Word dosyası başarıyla oluşturuldu: {filename}")

    def save_to_json(self, data, filename='academic_info.json'):
        """
        Verileri JSON dosyasına kaydeder
        Args:
            data: Kaydedilecek veriler
            filename: Kaydedilecek dosya adı
        """
        try:
            import json
            
            # Verileri JSON formatına uygun hale getir
            json_data = {
                'academic_info': {
                    'duties': data.get('academic_info', {}).get('duties', []),
                    'education': data.get('academic_info', {}).get('education', [])
                },
                'books': data.get('books', []),
                'articles': data.get('articles', []),
                'proceedings': data.get('proceedings', []),
                'projects': data.get('projects', []),
                'lessons': data.get('lessons', {})
            }
            
            # Dosyayı çalışılan dizine kaydet
            current_dir = os.getcwd()
            file_path = os.path.join(current_dir, filename)
            
            # JSON dosyasına kaydet (Türkçe karakterleri düzgün göstermek için ensure_ascii=False)
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, ensure_ascii=False, indent=4)
                
            print(f"JSON dosyası başarıyla oluşturuldu: {file_path}")
        except Exception as e:
            print(f"JSON dosyası kaydedilemedi: {e}")  

    def get_projects(self):
        """
        Akademisyenin projelerini çeker
        Returns:
            list: Proje bilgilerini içeren liste
        """
        try:
            # Projeler sekmesine tıkla
            projects_tab = self.wait.until(EC.element_to_be_clickable(
                (By.CSS_SELECTOR, "li#projectMenu a")))
            projects_tab.click()
            
            # Projelerin yüklenmesini bekle
            self.wait.until(EC.presence_of_element_located(
                (By.CSS_SELECTOR, "div.projectmain")))
            
            projects = []
            
            # Proje elementlerini bul
            project_elements = self.driver.find_elements(
                By.CSS_SELECTOR, 
                "div.projectmain"
            )
            
            pbar = tqdm(project_elements, desc="Projeler çekiliyor", unit="proje", colour="green")
            
            for project in pbar:
                try:
                    # Proje başlığı
                    title = project.find_element(By.CSS_SELECTOR, "span.baslika strong").text.strip()
                    
                    # Proje katılımcıları
                    contributors = []
                    contributor_elements = project.find_elements(By.CSS_SELECTOR, "a.popoverData")
                    for contributor in contributor_elements:
                        contributors.append(contributor.text.strip())
                    
                    # Proje detayları
                    project_type_div = project.find_element(By.CSS_SELECTOR, "div.projectType")
                    type_text = project_type_div.text.strip()
                    
                    # Detayları parçala
                    details = type_text.split(',')
                    
                    # Kurum ve tür bilgisi
                    institution_spans = project_type_div.find_elements(By.CSS_SELECTOR, "span.label")
                    institution = institution_spans[0].text.strip() if len(institution_spans) > 0 else ""
                    project_type = institution_spans[1].text.strip() if len(institution_spans) > 1 else ""
                    status = institution_spans[2].text.strip() if len(institution_spans) > 2 else ""
                    
                    # Tarih bilgisi
                    dates = details[1].strip() if len(details) > 1 else ""
                    
                    # Bütçe bilgisi
                    budget = f"{details[2].strip()} {details[3].strip()}" if len(details) > 3 else "Belirtilmemiş"
                    
                    project_info = {
                        'title': title,
                        'contributors': ', '.join(contributors),
                        'institution': institution,
                        'project_type': project_type,
                        'status': status,
                        'dates': dates,
                        'budget': budget
                    }
                    
                    projects.append(project_info)
                    pbar.set_description(f"Proje işleniyor: {title[:30]}...")
                    
                except Exception as e:
                    print(Fore.RED + f"HATA: Proje bilgisi çekilemedi: {e}")
                    continue
                
            return projects
            
        except Exception as e:
            print(Fore.RED + f"Projeler çekilirken hata: {e}")
            return []  

    def get_lessons(self):
        try:
            # Dersler sekmesine tıkla
            lessons_tab = self.wait.until(EC.element_to_be_clickable(
                (By.CSS_SELECTOR, "li#lessonMenu a")))
            lessons_tab.click()
            time.sleep(1)  # Menünün yüklenmesi için bekle
            
            # Önce "Önlisans" başlığına tıkla
            onlisans_header = self.wait.until(EC.element_to_be_clickable(
                (By.XPATH, "//a[contains(text(), 'Önlisans')]")))
            onlisans_header.click()
            time.sleep(1)  # Tablonun yüklenmesi için bekle
            
            # Eğitim seviyelerini tanımla
            education_levels = {
                'associate': {'id': 'collapse0', 'name': 'Ön Lisans'},
                'undergraduate': {'id': 'collapse1', 'name': 'Lisans'},
                'graduate': {'id': 'collapse2', 'name': 'Yüksek Lisans'}
            }
            
            all_lessons = {}
            
            # Her eğitim seviyesi için dersleri çek
            for level, info in education_levels.items():
                try:
                    # Eğitim seviyesi başlığına tıkla
                    level_header = self.wait.until(EC.element_to_be_clickable(
                        (By.XPATH, f"//a[contains(@href, '#{info['id']}')]")))
                    level_header.click()
                    time.sleep(1)  # Tablonun yüklenmesi için bekle
                    
                    lessons = []
                    
                    # Tablo elementini bul
                    table = self.wait.until(EC.presence_of_element_located(
                        (By.CSS_SELECTOR, f"#{info['id']} table.table")))
                    
                    # Tablo satırlarını al
                    rows = table.find_elements(By.CSS_SELECTOR, "tbody tr")
                    
                    for row in rows:
                        try:
                            cols = row.find_elements(By.TAG_NAME, "td")
                            if len(cols) >= 4:
                                lesson = {
                                    'term': cols[0].text.strip(),
                                    'name': cols[1].text.strip(),
                                    'language': cols[2].text.strip(),
                                    'hours': cols[3].text.strip()
                                }
                                lessons.append(lesson)
                        except Exception as e:
                            print(Fore.RED + f"HATA: Ders satırı çekilemedi: {e}")
                            continue
                    
                    all_lessons[level] = {
                        'name': info['name'],
                        'lessons': lessons
                    }
                    
                except Exception as e:
                    print(Fore.RED + f"HATA: {info['name']} dersleri çekilemedi: {e}")
                    all_lessons[level] = {
                        'name': info['name'],
                        'lessons': []
                    }
            
            return all_lessons
            
        except Exception as e:
            print(Fore.RED + f"Dersler çekilirken hata: {e}")
            return {}  
